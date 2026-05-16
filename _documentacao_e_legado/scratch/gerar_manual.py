from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

BASE = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(BASE, '..', 'MANUAL_UTILIZADOR_OAGB.md')
OUT_PATH = os.path.join(BASE, '..', 'MANUAL_UTILIZADOR_OAGB.docx')

doc = Document()

# --- PAGE MARGINS ---
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# --- STYLES ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_heading_style(para, level, text):
    sizes   = {1: 20, 2: 16, 3: 13, 4: 12}
    colors  = {1: '1B3A6B', 2: '1B3A6B', 3: '2E5FA3', 4: '333333'}
    bold_lv = {1: True, 2: True, 3: True, 4: True}
    para.style = f'Heading {level}'
    run = para.runs[0] if para.runs else para.add_run(text)
    if not para.runs:
        para.clear()
        run = para.add_run(text)
    run.font.size  = Pt(sizes.get(level, 12))
    run.font.bold  = bold_lv.get(level, False)
    run.font.color.rgb = RGBColor.from_string(colors.get(level, '000000'))
    para.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    para.paragraph_format.space_after  = Pt(4)

def add_table_from_lines(doc, lines):
    rows = [l.strip('|').split('|') for l in lines if l.strip().startswith('|') and '---' not in l]
    if len(rows) < 2:
        return
    cols = len(rows[0])
    tbl  = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            cell_text = cell.strip()
            # strip markdown bold
            cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
            cell_text = re.sub(r'`(.*?)`', r'\1', cell_text)
            tc = tbl.cell(ri, ci)
            tc.text = cell_text
            para = tc.paragraphs[0]
            run  = para.runs[0] if para.runs else para.add_run(cell_text)
            run.font.size = Pt(9.5)
            if ri == 0:
                run.font.bold = True
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), '1B3A6B')
                tc._tc.get_or_add_tcPr().append(shading)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()

def add_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('ORDEM DOS ADVOGADOS DA GUINÉ-BISSAU')
    r.font.size  = Pt(18)
    r.font.bold  = True
    r.font.color.rgb = RGBColor.from_string('1B3A6B')

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run('Manual de Utilizador e Guia de Gestão')
    r2.font.size = Pt(15)
    r2.font.bold = True

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run('Portal Institucional — oagb.gw')
    r3.font.size  = Pt(13)
    r3.font.color.rgb = RGBColor.from_string('2E5FA3')

    doc.add_paragraph()
    t4 = doc.add_paragraph()
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = t4.add_run('Versão 1.0  |  Maio 2026')
    r4.font.size  = Pt(11)
    r4.font.color.rgb = RGBColor.from_string('666666')

    doc.add_page_break()

add_cover(doc)

with open(MD_PATH, encoding='utf-8') as f:
    lines = f.readlines()

i = 0
table_buf = []

def flush_table(doc, buf):
    if buf:
        add_table_from_lines(doc, buf)
    return []

def add_inline_text(para, text):
    """Parse **bold**, `code`, ✅ etc. inline."""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.font.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
        else:
            para.add_run(part)

while i < len(lines):
    line = lines[i].rstrip('\n').rstrip('\r')

    # Skip YAML front-matter / horizontal rules used as decorators
    if line.strip() in ('---', ''):
        if table_buf:
            table_buf = flush_table(doc, table_buf)
        i += 1
        continue

    # TABLE rows
    if line.strip().startswith('|'):
        table_buf.append(line)
        i += 1
        continue
    else:
        if table_buf:
            table_buf = flush_table(doc, table_buf)

    # HEADINGS
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level = len(m.group(1))
        text  = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', m.group(2))  # strip links
        para  = doc.add_paragraph()
        para.add_run(text)
        set_heading_style(para, level, text)
        i += 1
        continue

    # BLOCKQUOTE (> ...)
    if line.strip().startswith('>'):
        text = line.strip().lstrip('> ').strip()
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        para = doc.add_paragraph(style='Quote' if 'Quote' in [s.name for s in doc.styles] else 'Normal')
        para.paragraph_format.left_indent = Cm(0.8)
        r = para.add_run(text)
        r.font.italic = True
        r.font.color.rgb = RGBColor.from_string('666666')
        i += 1
        continue

    # BULLET LIST (- or *)
    m_li = re.match(r'^(\s*)[-*]\s+(.*)', line)
    if m_li:
        indent = len(m_li.group(1)) // 2
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
        add_inline_text(para, m_li.group(2))
        i += 1
        continue

    # NUMBERED LIST
    m_nl = re.match(r'^\d+\.\s+(.*)', line)
    if m_nl:
        para = doc.add_paragraph(style='List Number')
        add_inline_text(para, m_nl.group(1))
        i += 1
        continue

    # NORMAL PARAGRAPH
    if line.strip():
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        add_inline_text(para, line.strip())

    i += 1

if table_buf:
    flush_table(doc, table_buf)

doc.save(OUT_PATH)
print(f"DOCX gerado: {OUT_PATH}")
